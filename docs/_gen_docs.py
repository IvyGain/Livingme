"""
Living Me ドキュメント生成スクリプト
requirements.docx と screen-flow.docx を生成する
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ──────────────────────────────────────────
# ユーティリティ
# ──────────────────────────────────────────

BRAND_BROWN  = RGBColor(0x6B, 0x4F, 0x3A)  # #6B4F3A
BRAND_ACCENT = RGBColor(0xC0, 0x70, 0x52)  # #C07052
LIGHT_GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
MID_GRAY     = RGBColor(0x88, 0x88, 0x88)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x00, 0x00, 0x00)
DARK         = RGBColor(0x33, 0x33, 0x33)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    # RGBColor は (r, g, b) のタプルとしてインデックスアクセス
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color or BRAND_BROWN
        run.bold = True
    return p


def add_body(doc, text, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_table(doc, headers, rows, header_bg=None, stripe=True):
    """ヘッダー付きテーブルを追加する"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ヘッダー行
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(hdr_cells[i], header_bg or BRAND_BROWN)

    # データ行
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)
            if stripe and r_idx % 2 == 1:
                set_cell_bg(row_cells[c_idx], RGBColor(0xFA, 0xF8, 0xF5))
    return table


def add_box(doc, title, lines, border_color=None):
    """枠付きのメモボックスを追加する"""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"【{title}】")
    run.bold = True
    run.font.color.rgb = border_color or BRAND_ACCENT
    run.font.size = Pt(10)
    for line in lines:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(line).font.size = Pt(9.5)
    doc.add_paragraph()


def add_flow_text(doc, steps):
    """テキストによる処理フロー（ステップ番号付き）"""
    for i, step in enumerate(steps):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        run.font.size = Pt(10)


def add_tree(doc, items):
    """インデント付きツリー表示"""
    for indent, text in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent * 0.8)
        prefix = "├─ " if indent > 0 else "▶ "
        run = p.add_run(prefix + text)
        run.font.size = Pt(9.5)
        if indent == 0:
            run.bold = True
            run.font.color.rgb = BRAND_BROWN


# ──────────────────────────────────────────
# 1. 要件定義書
# ──────────────────────────────────────────

def build_requirements():
    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # デフォルトフォント
    doc.styles['Normal'].font.name = 'メイリオ'
    doc.styles['Normal'].font.size = Pt(10.5)

    # ─── 表紙ブロック ───
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Living Me")
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_BROWN
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("〜 本当の自分と出逢う会 〜")
    run2.font.size = Pt(14)
    run2.font.color.rgb = BRAND_ACCENT
    run2.italic = True

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("要　件　定　義　書")
    run3.font.size = Pt(18)
    run3.bold = True
    run3.font.color.rgb = DARK

    doc.add_paragraph()

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Table Grid'
    meta = [
        ("バージョン", "v2.0"),
        ("最終更新日", "2026年3月30日"),
        ("対象読者", "初心者・非エンジニアの方"),
        ("作成者", "Living Me 運営チーム"),
    ]
    for i, (k, v) in enumerate(meta):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(meta_table.rows[i].cells[0], RGBColor(0xF0, 0xEB, 0xE5))
        meta_table.rows[i].cells[1].text = v
        meta_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ─── 1. サービス概要 ───
    add_heading(doc, "1. サービス概要", level=1)

    add_body(doc, "「Living Me」は、心の居場所づくりをテーマにした会員制コミュニティポータルサイトです。"
                  "動画アーカイブ・イベント・チャット・ジャーナリングなどの機能を通じて、"
                  "仲間との繋がりと自己表現をサポートします。")

    doc.add_paragraph()
    add_heading(doc, "1.1 基本情報", level=2)
    add_table(doc,
        ["項目", "内容"],
        [
            ("サービス名", "Living Me 〜本当の自分と出逢う会〜"),
            ("コンセプト", "愛のある心の居場所 / 頑張るよりも整える / 変わるよりも思い出す"),
            ("サイト種別", "会員向けコミュニティポータルサイト"),
            ("想定会員数（初期）", "40名"),
            ("想定会員数（半年後）", "150名"),
            ("想定会員数（将来）", "3,000〜5,000名"),
        ]
    )

    # ─── 2. ユーザー種別 ───
    doc.add_paragraph()
    add_heading(doc, "2. ユーザー種別と権限", level=1)

    add_body(doc, "サイトを利用するユーザーは以下の5種類に分かれます。それぞれできることが異なります。")
    doc.add_paragraph()

    add_heading(doc, "2.1 会員ロール一覧", level=2)
    add_table(doc,
        ["種別", "月額料金", "説明", "アクセス範囲"],
        [
            ("一般会員（MEMBER）", "3,300円/月", "標準の有料会員", "会員ページ全般"),
            ("紹介アンバサダー（REFERRAL）", "4,300円/月", "紹介活動を行うアンバサダー会員", "会員ページ + アンバサダー限定コンテンツ"),
            ("提携アンバサダー（PARTNER）", "9,900円/月", "提携事業者のアンバサダー会員", "会員ページ + アンバサダー限定コンテンツ"),
            ("管理者（ADMIN）", "—", "サイトを管理・運営するスタッフ", "管理画面 + 会員ページ全般"),
            ("未ログインユーザー", "—", "まだ会員登録していない訪問者", "LP・お問い合わせのみ"),
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "2.2 アカウント状態", level=2)
    add_table(doc,
        ["状態", "説明", "影響"],
        [
            ("有効（isActive = true）", "決済が正常で会員として有効", "すべての会員ページにアクセス可"),
            ("停止中（isActive = false）", "決済停止・未払い等でアカウント停止", "ログインできない。停止中の案内ページへ"),
        ]
    )

    # ─── 3. 機能一覧 ───
    doc.add_paragraph()
    add_heading(doc, "3. 機能一覧", level=1)

    add_heading(doc, "3.1 会員向け機能", level=2)
    add_table(doc,
        ["No.", "機能名", "説明"],
        [
            ("1", "ホーム画面", "ログイン後の最初の画面。今日のテーマ・ジャーナリング・アーカイブ・イベント・チャットを表示"),
            ("2", "アーカイブ動画", "過去の朝会・夜会・イベントの録画動画をいつでも視聴可能。カテゴリ・タグ検索対応"),
            ("3", "イベント", "朝会・夜会・ワークショップ等のイベント一覧。Zoomリンク確認・申込ができる"),
            ("4", "ジャーナリング", "毎日のテーマに沿って自分の内側を書き留める日記機能"),
            ("5", "チャット", "チャンネル別のリアルタイムチャット。スレッド（返信）にも対応"),
            ("6", "申請フォーム", "マヤ暦講座・個人セッション等の申込フォーム"),
            ("7", "アンバサダーページ", "自分の紹介リンクの確認・アンバサダー情報の表示"),
            ("8", "会員設定", "プロフィール編集・パスワード変更"),
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "3.2 管理者向け機能", level=2)
    add_table(doc,
        ["No.", "機能名", "説明"],
        [
            ("1", "ダッシュボード", "会員数・コンテンツ件数などのサマリー統計を表示"),
            ("2", "会員管理", "会員一覧・有効/無効の切替・アンバサダー種別の設定・紹介者の確認"),
            ("3", "紹介管理", "誰が誰を紹介したかの一覧・報酬計算結果の表示（単価は設定画面で変更可能）"),
            ("4", "問い合わせ管理", "一般ユーザーからの問い合わせ受信・返信（返信時にメール自動送信）・ステータス管理"),
            ("5", "コンテンツ管理（今日）", "ホーム画面に表示する今日のテーマ・ジャーナリングテーマの設定"),
            ("6", "コンテンツ管理（アーカイブ）", "動画アーカイブの作成・編集・公開設定"),
            ("7", "イベント管理", "イベントの作成・参加者一覧の確認"),
            ("8", "LP設定", "ランディングページの全セクションを編集（リアルタイムプレビュー付き）"),
            ("9", "ホーム画面設定", "ホームに表示するセクションの表示/非表示切替"),
            ("10", "チャット管理", "チャンネルの作成・ロール別アクセス権限の設定"),
            ("11", "外部サービス設定", "UnivaPay・Lark・紹介報酬単価などの設定値を管理画面から変更"),
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "3.3 LP（ランディングページ）設定の詳細", level=2)
    add_body(doc, "LPは管理画面からノーコードで全セクションを編集できます。")
    add_table(doc,
        ["セクション名", "編集できる内容"],
        [
            ("ファーストビュー（Hero）", "見出し・サブ見出し・背景色・背景画像・コンテンツ画像"),
            ("コンセプト（About）", "見出し・本文・コンセプト文言リスト（♥マーク付き）・画像"),
            ("お試し動画（Videos）", "YouTube埋め込みURL・Larkアーカイブ動画の選択・画像"),
            ("活動内容（Activities）", "各ブロックのタイトル・説明・画像（1件ずつ追加・削除可能）"),
            ("口コミ（Testimonials）", "投稿者名・肩書き・本文・アバター画像（1件ずつ追加・削除可能）"),
            ("最終CTA", "見出し・本文・「今すぐ始める」ボタンと「ログイン」ボタンの2種類のボタンテキスト"),
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "3.4 公開ページ（ログイン不要）", level=2)
    add_table(doc,
        ["ページ名", "URL", "説明"],
        [
            ("ランディングページ", "/", "サービス紹介・お試し動画・入会CTAボタン"),
            ("お問い合わせ", "/contact", "名前・メール・件名・内容を送信。管理者へ通知メールが届く"),
            ("会員登録", "/join", "UnivaPay決済フォーム。決済完了後に自動でアカウント作成"),
            ("ログイン", "/login", "メール+パスワードで認証"),
            ("招待登録", "/invite/[token]", "管理者が送った招待メールのURLからパスワードを設定して登録"),
            ("パスワードリセット", "/forgot-password", "メールアドレスを入力してリセットメールを受信"),
        ]
    )

    # ─── 4. 技術選定 ───
    doc.add_paragraph()
    add_heading(doc, "4. 技術選定", level=1)
    add_body(doc, "Living Me を構築するために使用している主要な技術とその選定理由を説明します。")
    doc.add_paragraph()

    add_heading(doc, "4.1 技術スタック一覧", level=2)
    add_table(doc,
        ["カテゴリ", "採用技術", "選定理由"],
        [
            ("フロントエンド\n＋バックエンド", "Next.js 15\n（App Router）",
             "フロントエンドとバックエンドを1つのプロジェクトで管理できる。世界中で広く使われており情報が豊富。Vercelとの相性が抜群"),
            ("プログラム言語", "TypeScript",
             "JavaScriptに型チェック機能を追加した言語。ミスを事前に検出でき、大規模開発でも安心"),
            ("デザイン", "Tailwind CSS",
             "クラス名を書くだけでスタイルが適用できるCSSフレームワーク。デザインの統一が簡単"),
            ("メインDB\n（機密情報）", "Neon（PostgreSQL）",
             "クラウド上のデータベースサービス。ユーザー情報・チャット・イベント等の大切なデータを安全に管理"),
            ("コンテンツDB\n（非機密）", "Lark Base（飛書）",
             "ノーコードでテーブルを管理できるサービス。動画アーカイブ・今日のテーマ等をエンジニア不要で更新可能"),
            ("認証", "NextAuth v5",
             "ログイン・セッション管理の実績あるライブラリ。JWTトークンで30日間セッションを保持"),
            ("決済", "UnivaPay",
             "日本向けサブスクリプション決済サービス。Webhook連携でリアルタイムに会員状態を自動更新"),
            ("画像保存", "Vercel Blob",
             "Vercelが提供するクラウドストレージ。画像をアップロードして公開URLで参照。管理者のみアップロード可"),
            ("メール送信", "Gmail SMTP\n（nodemailer）",
             "Gmailのアプリパスワードを使ってメールを送信。招待・ウェルカム・問い合わせ返信に使用"),
            ("ホスティング", "Vercel",
             "Next.jsの開発元が提供するホスティング。GitHubにコードをpushすると自動でサイトが更新される"),
        ],
        header_bg=BRAND_BROWN
    )

    doc.add_paragraph()
    add_heading(doc, "4.2 データの保存先の使い分け", level=2)
    add_body(doc, "データの種類によって保存先を使い分けています。")
    doc.add_paragraph()
    add_table(doc,
        ["保存するデータ", "保存先", "理由"],
        [
            ("会員情報・パスワード・決済記録", "Neon（PostgreSQL）", "機密性が高いため、厳格なアクセス管理ができるRDBに保存"),
            ("動画アーカイブ・今日のテーマ・イベント", "Lark Base（飛書）", "運営担当者がノーコードで更新できるようLarkに保存"),
            ("LPの設定・ホーム設定・外部サービス設定", "Neon（Setting テーブル）", "管理画面から変更できる設定値をDBに保存"),
            ("アップロード画像（LP・活動内容等）", "Vercel Blob", "クラウドストレージに保存し、公開URLで参照"),
        ]
    )

    # ─── 5. データモデル ───
    doc.add_paragraph()
    add_heading(doc, "5. データモデル（主要テーブル）", level=1)
    add_body(doc, "データベースには以下のテーブル（データの入れ物）があります。")
    doc.add_paragraph()
    add_table(doc,
        ["テーブル名", "説明", "主なデータ項目"],
        [
            ("User（ユーザー）", "会員情報を管理する中心テーブル",
             "メール・パスワード(暗号化)・名前・ロール・アンバサダー種別・紹介者・参加日"),
            ("ChatChannel（チャンネル）", "チャットのチャンネル情報",
             "チャンネル名・説明・必要ロール・アーカイブ状態"),
            ("ChatMessage（メッセージ）", "チャットのメッセージ",
             "内容・投稿者・チャンネル・親メッセージID（スレッド用）"),
            ("InviteToken（招待トークン）", "管理者が発行する招待リンク情報",
             "トークン文字列・宛先メール・有効期限（72時間）"),
            ("EventRegistration（イベント参加）", "イベント申込の記録",
             "イベントID・参加者ID・回答内容"),
            ("ContactInquiry（問い合わせ）", "お問い合わせフォームの受信内容",
             "名前・メール・件名・本文・ステータス（未対応/返信済/クローズ）"),
            ("InquiryReply（問い合わせ返信）", "問い合わせへの返信内容",
             "返信内容・メール送信済みフラグ"),
            ("Setting（設定値）", "管理画面から変更できる設定値",
             "キー・値・秘密情報フラグ（True=暗号化保存）"),
            ("UnivaPayEvent（決済ログ）", "UnivaPay Webhookの受信ログ",
             "イベントID・種類・受信日時・ペイロード（冪等性制御）"),
        ]
    )

    # ─── 6. セキュリティ ───
    doc.add_paragraph()
    add_heading(doc, "6. セキュリティ設計", level=1)
    add_table(doc,
        ["対策", "内容"],
        [
            ("パスワード保護", "bcrypt（ビークリプト）でハッシュ化して保存。平文（そのままの文字列）では保存しない"),
            ("ログイン試行制限", "5回連続失敗で15分間ブロック（ブルートフォース攻撃対策）"),
            ("セッション管理", "JWT（認証トークン）を使用。有効期間30日。ブラウザのCookieに保存"),
            ("画像アップロード制限", "管理者(ADMIN)のみアップロード可。JPEG/PNG/WebP/GIFのみ。最大10MB"),
            ("管理画面アクセス制限", "role=ADMIN かつ isActive=true の場合のみアクセス可"),
            ("機密設定値の暗号化", "UnivaPay・Larkのシークレットキーはデータベースに暗号化して保存"),
            ("Webhook冪等性", "UnivaPay Webhookは同じイベントが2回届いても1回だけ処理される"),
        ]
    )

    # ─── 7. 非機能要件 ───
    doc.add_paragraph()
    add_heading(doc, "7. 非機能要件", level=1)

    add_heading(doc, "7.1 パフォーマンス目標", level=2)
    add_table(doc,
        ["指標", "目標値", "説明"],
        [
            ("FCP（最初のコンテンツ表示）", "1.8秒以内", "ページを開いてから最初の表示が出るまでの時間"),
            ("LCP（最大コンテンツ表示）", "2.5秒以内", "最も大きなコンテンツが表示されるまでの時間"),
            ("CLS（レイアウトのずれ）", "0.1以下", "ページが読み込まれる際にコンテンツがずれる量"),
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "7.2 対応ブラウザ・デバイス", level=2)
    add_table(doc,
        ["区分", "対応内容"],
        [
            ("ブラウザ", "Chrome・Safari・Firefox・Edge（最新版）"),
            ("デバイス", "PC・スマートフォン・タブレット（レスポンシブ対応）"),
            ("アクセシビリティ", "WCAG 2.1 AA準拠を目標。コントラスト比4.5:1以上"),
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "7.3 スケール計画", level=2)
    add_table(doc,
        ["フェーズ", "会員数", "対応"],
        [
            ("現在", "〜150名", "Vercel Hobby/Pro + Neon無料プランで対応可能"),
            ("成長期", "〜1,000名", "Vercel Pro + Neon有料プランへ移行"),
            ("大規模", "1,000〜5,000名", "インフラのスケールアップを検討"),
        ]
    )

    # ─── 8. 用語集 ───
    doc.add_paragraph()
    add_heading(doc, "8. 用語集（初心者向け）", level=1)
    add_table(doc,
        ["用語", "意味"],
        [
            ("JWT（ジェイダブリューティー）", "ログイン情報を安全に保存するための「証明書」のようなもの。30日間有効"),
            ("Webhook（ウェブフック）", "外部サービス（UnivaPay等）からの「通知」。決済完了時などに自動で呼び出される"),
            ("bcrypt（ビークリプト）", "パスワードを安全な形式に変換する暗号化の仕組み"),
            ("ロール（role）", "ユーザーの役割・権限のこと。MEMBER（会員）やADMIN（管理者）など"),
            ("isActive", "アカウントが有効かどうかを示すフラグ。true=有効、false=停止中"),
            ("API（エーピーアイ）", "システム同士がデータをやり取りするための窓口"),
            ("CRUD", "Create（作成）・Read（読取）・Update（更新）・Delete（削除）の4操作の総称"),
            ("Mermaid（マーメイド）", "テキストで図を描くための記法。フローチャートなどを書ける"),
            ("PostgreSQL（ポストグレスキューエル）", "広く使われているオープンソースのデータベース管理システム"),
            ("Vercel（バーセル）", "Next.jsのホスティングサービス。GitHubと連携して自動デプロイできる"),
        ]
    )

    path = os.path.join(OUTPUT_DIR, "要件定義書.docx")
    doc.save(path)
    print(f"✓ 保存: {path}")


# ──────────────────────────────────────────
# 2. 画面遷移図
# ──────────────────────────────────────────

def build_screen_flow():
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    doc.styles['Normal'].font.name = 'メイリオ'
    doc.styles['Normal'].font.size = Pt(10.5)

    # ─── 表紙 ───
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Living Me")
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_BROWN
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("〜 本当の自分と出逢う会 〜")
    run2.font.size = Pt(14)
    run2.font.color.rgb = BRAND_ACCENT
    run2.italic = True

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("画　面　遷　移　図　／　サイトマップ")
    run3.font.size = Pt(18)
    run3.bold = True
    run3.font.color.rgb = DARK

    doc.add_paragraph()
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = 'Table Grid'
    for i, (k, v) in enumerate([("最終更新日", "2026年3月30日"), ("対象読者", "初心者・非エンジニアの方"), ("備考", "全ページのURL・説明・アクセス条件を掲載")]):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(meta_table.rows[i].cells[0], RGBColor(0xF0, 0xEB, 0xE5))
        meta_table.rows[i].cells[1].text = v
        meta_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ─── アクセス権限の凡例 ───
    add_heading(doc, "アクセス権限の凡例", level=1)
    add_table(doc,
        ["マーク", "意味", "対象ユーザー"],
        [
            ("🌐 公開", "ログイン不要でアクセスできる", "誰でも"),
            ("🔐 会員", "ログイン済み＋有効会員が必要", "有料会員（MEMBER / REFERRAL / PARTNER）"),
            ("🛡 管理者", "管理者権限が必要", "role=ADMIN のユーザーのみ"),
        ]
    )
    doc.add_paragraph()
    add_box(doc, "アクセス権限がない場合の動作",
        [
            "未ログイン状態で会員ページにアクセス → ログインページ（/login）にリダイレクト",
            "アカウント停止中でログイン → 停止中エラーページへ",
            "管理者でないユーザーが管理画面へアクセス → ホーム（/home）にリダイレクト",
        ]
    )

    # ─── サイトマップ（ツリー形式） ───
    add_heading(doc, "サイトマップ（全ページ一覧）", level=1)
    add_body(doc, "以下がサイト内の全ページです。インデントが深いほど下の階層のページです。")
    doc.add_paragraph()

    add_heading(doc, "🌐 公開ページ（ログイン不要）", level=2)
    add_tree(doc, [
        (0, "/  — ランディングページ（LP）"),
        (1, "サービスの紹介・お試し動画・入会ボタン"),
        (0, "/contact  — お問い合わせ"),
        (1, "名前・メール・件名・内容を送信"),
        (1, "送信後、管理者へ通知メールが届く"),
        (0, "/join  — 会員登録・決済"),
        (1, "UnivaPay決済フォーム"),
        (1, "決済完了後にアカウントが自動作成される"),
        (0, "/login  — ログイン"),
        (1, "メールアドレス＋パスワードで認証"),
        (0, "/invite/[token]  — 招待リンクからの登録"),
        (1, "管理者が送った招待メールのリンクからアクセス"),
        (1, "名前・パスワードを設定してアカウント作成"),
        (0, "/forgot-password  — パスワードリセット申請"),
        (0, "/reset-password  — パスワード再設定"),
        (0, "/demo  — デモページ（認証不要のサービス体験）"),
        (1, "/demo/archive  — アーカイブ（サンプル）"),
        (1, "/demo/events  — イベント（サンプル）"),
        (1, "/demo/journal  — ジャーナル（サンプル）"),
        (1, "/demo/admin  — 管理画面（サンプル）"),
    ])

    doc.add_paragraph()
    add_heading(doc, "🔐 会員ページ（ログイン済み・有料会員のみ）", level=2)
    add_tree(doc, [
        (0, "/home  — ホーム画面"),
        (1, "今日のテーマ・ジャーナリングテーマ"),
        (1, "最新アーカイブ動画・イベント・チャット"),
        (0, "/archive  — アーカイブ一覧"),
        (1, "カテゴリ・タグで検索"),
        (1, "/archive/[id]  — アーカイブ詳細（動画再生）"),
        (0, "/events  — イベント一覧"),
        (1, "朝会・夜会・ワークショップ等"),
        (1, "/events/[id]  — イベント詳細（Zoomリンク・申込）"),
        (0, "/journal  — ジャーナル一覧"),
        (1, "自分が書いた日記の一覧"),
        (1, "/journal/new  — ジャーナル作成"),
        (0, "/chat  — チャット（チャンネル一覧）"),
        (1, "/chat/[channelId]  — チャンネル内メッセージ"),
        (1, "スレッド（返信）機能あり"),
        (0, "/forms  — 申請フォーム一覧"),
        (1, "/forms/[slug]  — 申請フォーム詳細・入力"),
        (0, "/ambassador  — アンバサダーページ"),
        (1, "自分の紹介リンク・アンバサダー情報"),
        (0, "/about  — サービス概要"),
        (0, "/settings  — 会員設定（プロフィール・パスワード）"),
    ])

    doc.add_paragraph()
    add_heading(doc, "🛡 管理画面（管理者のみ）", level=2)
    add_tree(doc, [
        (0, "/admin  — 管理ダッシュボード"),
        (1, "会員数・コンテンツ件数のサマリー"),
        (0, "/admin/members  — 会員管理"),
        (1, "会員一覧・有効/無効切替"),
        (1, "アンバサダー種別の設定"),
        (1, "紹介者の確認"),
        (0, "/admin/referrals  — 紹介管理"),
        (1, "誰が誰を紹介したかの一覧"),
        (1, "報酬計算結果（単価×紹介件数）"),
        (0, "/admin/inquiries  — 問い合わせ管理"),
        (1, "受信一覧・ステータス管理"),
        (1, "返信するとメールが自動送信される"),
        (0, "/admin/content/today  — 今日のコンテンツ管理"),
        (1, "ホーム画面に表示するテーマを設定"),
        (0, "/admin/content/archives  — アーカイブ管理"),
        (1, "動画の作成・編集・公開/非公開切替"),
        (1, "/admin/content/archives/new  — アーカイブ作成"),
        (1, "/admin/content/archives/[id]/edit  — 編集"),
        (0, "/admin/events  — イベント管理"),
        (1, "/admin/events/[id]/registrations  — 参加者一覧"),
        (0, "/admin/lp-settings  — LP設定"),
        (1, "ランディングページの全セクション編集"),
        (1, "リアルタイムプレビュー付き"),
        (0, "/admin/home-layout  — ホーム画面設定"),
        (1, "表示するセクションの切替"),
        (0, "/admin/chat  — チャット管理"),
        (1, "チャンネルの作成・ロール別権限設定"),
        (0, "/admin/settings  — 外部サービス設定"),
        (1, "UnivaPay・Lark・紹介報酬単価の設定"),
    ])

    # ─── 全画面一覧テーブル ───
    doc.add_paragraph()
    add_heading(doc, "全画面一覧（詳細）", level=1)

    add_heading(doc, "公開ページ", level=2)
    add_table(doc,
        ["URL", "画面名", "説明", "アクセス"],
        [
            ("/", "ランディングページ（LP）", "サービス紹介・お試し動画・CTAボタン", "🌐 誰でも"),
            ("/contact", "お問い合わせ", "名前・メール・内容を送信。管理者に通知メール", "🌐 誰でも"),
            ("/join", "会員登録・決済", "UnivaPay決済フォーム", "🌐 誰でも"),
            ("/login", "ログイン", "メール＋パスワード認証", "🌐 誰でも"),
            ("/invite/[token]", "招待登録", "招待メールのURLからパスワード設定", "🌐 招待された人"),
            ("/forgot-password", "パスワードリセット申請", "メールを入力してリセットリンクを受信", "🌐 誰でも"),
            ("/reset-password", "パスワード再設定", "新しいパスワードを入力", "🌐 誰でも"),
            ("/demo", "デモホーム", "認証不要のサービス体験ページ", "🌐 誰でも"),
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "会員ページ", level=2)
    add_table(doc,
        ["URL", "画面名", "説明", "アクセス"],
        [
            ("/home", "ホーム", "今日のテーマ・アーカイブ・イベント・チャット", "🔐 有料会員"),
            ("/archive", "アーカイブ一覧", "過去動画・カテゴリ・タグ検索", "🔐 有料会員"),
            ("/archive/[id]", "アーカイブ詳細", "動画再生・説明文表示", "🔐 有料会員"),
            ("/events", "イベント一覧", "朝会・夜会・特別企画の一覧", "🔐 有料会員"),
            ("/events/[id]", "イベント詳細", "Zoomリンク確認・申込フォーム", "🔐 有料会員"),
            ("/journal", "ジャーナル一覧", "自分の日記一覧", "🔐 有料会員"),
            ("/journal/new", "ジャーナル作成", "今日のテーマに沿って書く", "🔐 有料会員"),
            ("/chat", "チャット（一覧）", "チャンネル選択画面", "🔐 有料会員"),
            ("/chat/[channelId]", "チャット（チャンネル内）", "メッセージ・スレッド表示", "🔐 有料会員"),
            ("/forms", "申請フォーム一覧", "マヤ暦講座・個人セッション等", "🔐 有料会員"),
            ("/forms/[slug]", "申請フォーム詳細", "申込フォーム入力・送信", "🔐 有料会員"),
            ("/ambassador", "アンバサダーページ", "紹介リンク確認・アンバサダー情報", "🔐 有料会員"),
            ("/about", "サービス概要", "Living Meについての説明", "🔐 有料会員"),
            ("/settings", "会員設定", "プロフィール編集・パスワード変更", "🔐 有料会員"),
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "管理画面", level=2)
    add_table(doc,
        ["URL", "画面名", "説明", "アクセス"],
        [
            ("/admin", "管理ダッシュボード", "会員数・コンテンツ統計のサマリー", "🛡 管理者"),
            ("/admin/members", "会員管理", "一覧・有効/無効・アンバサダー種別・紹介者確認", "🛡 管理者"),
            ("/admin/referrals", "紹介管理", "紹介関係の一覧・報酬計算結果", "🛡 管理者"),
            ("/admin/inquiries", "問い合わせ管理", "受信一覧・返信（メール自動送信）", "🛡 管理者"),
            ("/admin/content/today", "今日のコンテンツ", "ホームのテーマ・ジャーナリングテーマ設定", "🛡 管理者"),
            ("/admin/content/archives", "アーカイブ管理", "動画一覧・作成・編集・公開設定", "🛡 管理者"),
            ("/admin/content/archives/new", "アーカイブ作成", "新規動画の登録", "🛡 管理者"),
            ("/admin/content/archives/[id]/edit", "アーカイブ編集", "既存動画の編集", "🛡 管理者"),
            ("/admin/events", "イベント管理", "イベント作成・参加者確認", "🛡 管理者"),
            ("/admin/events/[id]/registrations", "参加者一覧", "イベント参加者の確認", "🛡 管理者"),
            ("/admin/lp-settings", "LP設定", "LPの全セクション編集（プレビュー付き）", "🛡 管理者"),
            ("/admin/home-layout", "ホーム画面設定", "表示セクションの切替", "🛡 管理者"),
            ("/admin/chat", "チャット管理", "チャンネル作成・ロール別権限設定", "🛡 管理者"),
            ("/admin/settings", "外部サービス設定", "UnivaPay・Lark・紹介報酬単価の設定", "🛡 管理者"),
        ]
    )

    # ─── 主要な動線 ───
    doc.add_paragraph()
    add_heading(doc, "主要なユーザー動線", level=1)

    add_heading(doc, "① 新規会員の入会フロー", level=2)
    add_flow_text(doc, [
        "LP（/）でサービスを知る",
        "「今すぐ始める」ボタンをクリック → /join へ移動",
        "UnivaPay決済フォームでクレジットカード情報を入力・決済完了",
        "アカウントが自動作成される（メール受信）",
        "/login でメール＋パスワードを入力してログイン",
        "/home（ホーム画面）へ自動遷移",
    ])

    doc.add_paragraph()
    add_heading(doc, "② 管理者による招待フロー", level=2)
    add_flow_text(doc, [
        "管理者が /admin/members を開く",
        "「招待メールを送る」ボタンをクリック",
        "招待したい人のメールアドレスを入力して送信",
        "招待された人のメールに招待リンクが届く（72時間有効）",
        "招待リンクをクリック → /invite/[token] へ移動",
        "名前・パスワードを設定してアカウント作成完了",
        "/login でログイン → /home へ",
    ])

    doc.add_paragraph()
    add_heading(doc, "③ 一般ユーザーからの問い合わせフロー", level=2)
    add_flow_text(doc, [
        "LP（/）フッターの「お問い合わせ」リンクをクリック",
        "/contact の問い合わせフォームに名前・メール・内容を入力して送信",
        "管理者のGmailに通知メールが届く",
        "管理者が /admin/inquiries で問い合わせ内容を確認",
        "返信フォームに返信内容を入力して「返信を送る」クリック",
        "問い合わせ者のメールアドレスに返信メールが自動送信される",
    ])

    doc.add_paragraph()
    add_heading(doc, "④ 管理者によるLP更新フロー", level=2)
    add_flow_text(doc, [
        "/admin/lp-settings を開く",
        "各セクションの「編集」ボタンをクリックして内容を変更",
        "画像は「ファイルを選択」ボタンでパソコンから直接アップロード",
        "右側のプレビューでリアルタイムに変更を確認",
        "「変更を保存する」ボタンをクリック",
        "LP（/）に変更が即時反映される",
    ])

    path = os.path.join(OUTPUT_DIR, "画面遷移図.docx")
    doc.save(path)
    print(f"✓ 保存: {path}")


# ──────────────────────────────────────────
# 実行
# ──────────────────────────────────────────

if __name__ == "__main__":
    build_requirements()
    build_screen_flow()
    print("完了！")
